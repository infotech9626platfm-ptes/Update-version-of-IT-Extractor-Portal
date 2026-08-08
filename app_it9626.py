#********** This script containing the use of Service account ID and the JSON KEY in TOML ********
import io
import os
import re
import fitz  # PyMuPDF
import streamlit as st
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Google API Libraries (Using Service Account Credentials)
from google.oauth2.service_account import Credentials
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload, MediaIoBaseUpload

# ==========================================
# 1. CONFIGURATION & DRIVE FOLDER MAPPING
# ==========================================
SYLLABUS_CODE = "9626"

# Google Drive Folder IDs mapped to your live Google Drive folders
FOLDER_IDS = {
    "theory": "1T1sIqRKxF5aO_r0sCyIVxidt0TyXOCcB",     # Theory Papers (P1 & P3)
    "practical": "1EWBiwjvTc12LVtyNi2V9P9RSr8d2vgq7",  # Practical Papers (P2 & P4)
    "zips": "1AsXq8TktyqajB7XTa9SQ5f85Pr6CQcFJ"          # Source Files (.zip)
}

# Local directories for mirroring files locally on the server
LOCAL_FOLDERS = {
    "theory": "9626_theory",
    "practical": "9626_practical",
    "zips": "9626_zips"
}

# Ensure local storage directories exist on server startup
for folder_path in LOCAL_FOLDERS.values():
    if not os.path.exists(folder_path):
        os.makedirs(folder_path)


# ==========================================
# 2. AUTOMATIC ROUTING & GOOGLE DRIVE API
# ==========================================
def determine_target_folder(filename: str) -> tuple[str, str]:
    """
    Analyzes the filename using Regular Expressions to determine target folder.
    Returns a tuple of (folder_key, folder_display_name).
    """
    filename_lower = filename.lower()
    
    # 1. Zip / Source Files
    if filename_lower.endswith(".zip") or "_sf_" in filename_lower:
        return "zips", "9626_zips (Source Files)"
        
    # 2. Practical Papers (Papers 02 and 04)
    if re.search(r'_(qp|ms)_0[24]\b', filename_lower):
        return "practical", "9626_practical (Papers 2 & 4)"
        
    # 3. Theory Papers (Papers 11, 12, 13, 31, 32, 33)
    if re.search(r'_(qp|ms)_(1[123]|3[123])\b', filename_lower):
        return "theory", "9626_theory (Papers 1 & 3)"
        
    return None, None


def build_drive_service():
    """
    Authenticates with Google Drive API using Service Account Credentials stored in Streamlit Secrets.
    """
    try:
        # Full access scope required for searching, downloading, and uploading files
        SCOPES = ['https://www.googleapis.com/auth/drive']
        
        # Parse service account secrets from Streamlit secrets dictionary
        service_account_info = dict(st.secrets["gcp_service_account"])
        
        # Build Service Account Credentials instance
        creds = Credentials.from_service_account_info(
            service_account_info, 
            scopes=SCOPES
        )
        return build('drive', 'v3', credentials=creds)
    except Exception as e:
        st.error(f"❌ Authentication Configuration Error: {e}")
        return None


def upload_file_to_drive(file_bytes, filename, folder_id, mime_type):
    """
    Uploads file binary stream directly to Google Drive using Service Account Credentials.
    """
    service = build_drive_service()
    if not service:
        return None

    try:
        file_stream = io.BytesIO(file_bytes)
        
        file_metadata = {
            'name': filename,
            'parents': [folder_id]
        }
        
        media = MediaIoBaseUpload(
            file_stream, 
            mimetype=mime_type, 
            resumable=True
        )

        uploaded_file = service.files().create(
            body=file_metadata,
            media_body=media,
            fields='id, webViewLink'
        ).execute()

        return uploaded_file

    except Exception as error:
        st.error(f"❌ Drive API Upload Failed: {error}")
        return None


def sync_drive_folder_to_local(folder_key: str) -> tuple[int, str]:
    """
    Queries Google Drive for a specific folder and downloads any files missing locally.
    Returns (downloaded_count, status_message).
    """
    service = build_drive_service()
    if not service:
        return 0, "Failed to authenticate with Google Drive."

    drive_folder_id = FOLDER_IDS[folder_key]
    local_path = LOCAL_FOLDERS[folder_key]

    try:
        query = f"'{drive_folder_id}' in parents and trashed = false"
        results = service.files().list(q=query, fields="files(id, name, mimeType)").execute()
        drive_files = results.get('files', [])

        downloaded_count = 0

        for file_info in drive_files:
            file_name = file_info['name']
            file_id = file_info['id']
            local_file_path = os.path.join(local_path, file_name)

            if not os.path.exists(local_file_path):
                request = service.files().get_media(fileId=file_id)
                with open(local_file_path, "wb") as f:
                    downloader = MediaIoBaseDownload(f, request)
                    done = False
                    while not done:
                        status, done = downloader.next_chunk()
                
                downloaded_count += 1

        return downloaded_count, f"Synced {downloaded_count} new file(s) for folder `{folder_key}`."

    except Exception as e:
        return 0, f"Sync error on folder `{folder_key}`: {e}"


# ==========================================
# 3. ADVANCED FLEXIBLE SEARCH ENGINE
# ==========================================
def search_pdfs(keyword_list, folder_path, allowed_variants):
    """
    Scans local PDF files for keywords using flexible, case-insensitive matching.
    """
    results = []
    if not os.path.exists(folder_path):
        return results

    cleaned_keywords = [k.strip().lower() for k in keyword_list if k.strip()]
    if not cleaned_keywords:
        return results

    for file in os.listdir(folder_path):
        if file.endswith(".pdf"):
            base_name = os.path.splitext(file)[0]
            
            is_valid_variant = any(base_name.endswith(f"_{variant}") for variant in allowed_variants)
            if not is_valid_variant:
                continue

            filepath = os.path.join(folder_path, file)
            try:
                doc = fitz.open(filepath)
                for page_num in range(len(doc)):
                    page_text = doc[page_num].get_text()
                    
                    matches_all = True
                    for kw in cleaned_keywords:
                        escaped_kw = re.escape(kw)
                        pattern = r'\b' + escaped_kw + r'(s|es)?\b'
                        
                        if not re.search(pattern, page_text, re.IGNORECASE) and kw not in page_text.lower():
                            matches_all = False
                            break
                    
                    if matches_all:
                        results.append({
                            "file": file,
                            "page": page_num,
                            "path": filepath,
                            "type": "QP" if "_qp_" in file else "MS"
                        })
                doc.close()
            except Exception:
                continue
                
    return results


def render_pdf_page_image(file_path: str, page_num: int) -> bytes:
    """
    Renders a specific page of a PDF file to PNG byte stream.
    """
    pdf_doc = fitz.open(file_path)
    page = pdf_doc.load_page(page_num)
    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # Crisp rendering
    img_bytes = pix.tobytes("png")
    pdf_doc.close()
    return img_bytes


# ==========================================
# 4. WORD DOCUMENT BUILDER HELPER
# ==========================================
def add_page_number_to_header(run):
    """
    Inserts a dynamic Word XML page number field into a header text run.
    """
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')

    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)


def create_custom_word_handout(basket_items, syllabus_code):
    """
    Generates a Word document with custom page and margin settings.
    """
    doc = Document()

    # Configure Section Page Dimensions & Margins
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.5)

        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.3)
        section.left_margin = Inches(0.4)
        section.right_margin = Inches(0.4)

        # Header Page Numbering (Top Center)
        header = section.header
        header_para = header.paragraphs[0]
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        header_run = header_para.add_run("Page ")
        header_run.font.name = "Arial"
        header_run.font.size = Pt(10)
        add_page_number_to_header(header_run)

    # Main Document Header
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(f'PTES {syllabus_code} IT Handout Worksheets')
    title_run.font.bold = True
    title_run.font.size = Pt(14)

    # Process saved pages
    for idx, item in enumerate(basket_items):
        h = doc.add_heading(f"Source: {item['file']} (Page {item['page'] + 1})", level=2)
        h.paragraph_format.space_before = Pt(4)
        h.paragraph_format.space_after = Pt(4)

        img_data = io.BytesIO(render_pdf_page_image(item['path'], item['page']))

        # Single image insert with proportional scaling (Width = 6.5")
        doc.add_picture(img_data, width=Inches(6.5))

        if idx < len(basket_items) - 1:
            doc.add_page_break()

    return doc


# ==========================================
# 5. APP STATE INITIALIZATION & AUTO-SYNC
# ==========================================
if 'handout_basket' not in st.session_state:
    st.session_state.handout_basket = []
if 'theory_results' not in st.session_state:
    st.session_state.theory_results = []
if 'practical_results' not in st.session_state:
    st.session_state.practical_results = []

# --- AUTO-SYNC PROCEDURE ON PORTAL WAKE UP / STARTUP ---
if 'has_auto_synced' not in st.session_state:
    st.session_state.has_auto_synced = False

if not st.session_state.has_auto_synced:
    with st.spinner("⚡ Portal waking up: Syncing latest files from Google Drive..."):
        total_auto_synced = 0
        for f_key in ["theory", "practical", "zips"]:
            count, _ = sync_drive_folder_to_local(f_key)
            total_auto_synced += count
        
        # Mark auto-sync complete so it only executes once per user session
        st.session_state.has_auto_synced = True
        
        if total_auto_synced > 0:
            st.toast(f"🔄 Auto-Sync Complete: Downloaded {total_auto_synced} new file(s)!")


# ==========================================
# 6. STREAMLIT UI LAYOUT & STYLING
# ==========================================
st.set_page_config(page_title="9626 IT Resource Platform", layout="wide")

MAIN_BG_COLOR = "#f9f0ee"     # Main screen background color
SIDEBAR_BG_COLOR = "#FFFDD0"  # Sidebar background color
INPUT_BG_COLOR = "#FA8FEB"    # Input box background color (Pink)
INPUT_BORDER_COLOR = "#1A1A1A" # Dark border color

st.markdown(
    f"""
    <style>
    /* Main application background */
    .stAppViewContainer {{
        background-color: {MAIN_BG_COLOR} !important;
    }}
    
    /* Top sticky header background */
    .stHeader {{
        background-color: {MAIN_BG_COLOR} !important;
    }}

    /* Sidebar background color */
    [data-testid="stSidebar"] {{
        background-color: {SIDEBAR_BG_COLOR} !important;
    }}

    /* Text Inputs & Selectboxes */
    div[data-baseweb="input"], 
    div[data-baseweb="base-input"],
    .stTextInput input, 
    .stPasswordInput input {{
        background-color: {INPUT_BG_COLOR} !important;
        border: 2px solid {INPUT_BORDER_COLOR} !important;
        border-radius: 8px !important;
        color: #000000 !important;
        font-weight: 600 !important;
    }}

    div[data-baseweb="input"] > input {{
        background-color: transparent !important;
    }}

    .stSelectbox div[data-baseweb="select"],
    div[data-baseweb="select"] > div {{
        background-color: {INPUT_BG_COLOR} !important;
        border: 5px solid {INPUT_BORDER_COLOR} !important;
        border-radius: 10px !important;
    }}

    .stSelectbox span,
    div[data-baseweb="select"] span {{
        color: #6D3761 !important;
        font-weight: 600 !important;
    }}

    .stSelectbox svg,
    div[data-baseweb="select"] svg {{
        fill: #6D3761 !important;
    }}

    /* Custom expander styling */
    .stExpander {{
        background-color: #ffffff;
        border-radius: 10px;
        border: 1px solid #e0e0e0;
        margin-bottom: 8px;
    }}
    </style>
    """,
    unsafe_allow_html=True
)

st.title("BRUNEI FORM SIXTH CENTRE")
st.subheader("💻 9626 Information Technology PYP Resources")

# ==========================================
# SIDEBAR: BASKET SUMMARY & MANUAL DRIVE SYNC
# ==========================================
with st.sidebar:
    st.header("🛒 Handout Basket Summary")
    st.metric(label="Saved Pages in Basket", value=len(st.session_state.handout_basket))
    
    if st.button("🗑️ Clear Basket", key="sb_clear_basket"):
        st.session_state.handout_basket = []
        st.rerun()

    st.markdown("---")
    st.header("🔄 Google Drive Sync")
    st.caption("Sync locally mirrored files with Google Drive.")
    
    # MANUAL SYNC PROCEDURE (TRIGGERED VIA BUTTON CLICK)
    if st.button("🔄 Sync All Files from Google Drive", type="primary", key="sb_sync_btn"):
        with st.spinner("Scanning Google Drive folders and downloading new files..."):
            total_synced = 0
            for f_key in ["theory", "practical", "zips"]:
                count, msg = sync_drive_folder_to_local(f_key)
                total_synced += count
                st.info(msg)
            
            st.success(f"🎉 Sync Complete! **{total_synced}** new file(s) downloaded!")

# Application Navigation Tabs
tab1, tab2, tab3, tab4, tab5 = st.tabs([
    "🔍 Theory Search(P1&P3)", 
    "⚙️ Practical Search(P2&P4)", 
    "🛒 Handout Cart", 
    "📦 Source Files(ZIP)", 
    "🔒 Admin Panel"
])


# --- TAB 1: THEORY SEARCH (P1 & P3) ---
with tab1:
    st.header("Search Theory Papers (Paper 1 & Paper 3)")
    st.caption("Variants: Paper 1 (11, 12, 13) | Paper 3 (31, 32, 33)")
    keyword_t1 = st.text_input("Enter Theory Keywords (e.g., 'Normalized', 'Relational Database', 'CSS')", key="t1_kw")

    if st.button("Search Theory Papers", type="primary"):
        if keyword_t1:
            with st.spinner("Scanning Theory PDFs..."):
                keywords = [k.strip() for k in keyword_t1.split(",") if k.strip()]
                theory_variants = ["11", "12", "13", "31", "32", "33"]
                st.session_state.theory_results = search_pdfs(keywords, LOCAL_FOLDERS["theory"], theory_variants)
        else:
            st.warning("Please enter a keyword.")

    if st.session_state.theory_results:
        st.write(f"Found **{len(st.session_state.theory_results)}** matching pages:")
        for idx, item in enumerate(st.session_state.theory_results):
            doc_kind = "📝 Question Paper" if item["type"] == "QP" else "🔑 Marking Scheme"
            label = f"📄 {item['file']} | {doc_kind} | Page {item['page'] + 1}"
            
            with st.expander(label):
                col_img, col_actions = st.columns([3, 1])
                
                with col_img:
                    img_data = render_pdf_page_image(item['path'], item['page'])
                    st.image(img_data, use_container_width=True)

                with col_actions:
                    st.write("### Actions")
                    if st.button("➕ Add to Basket", key=f"add_t1_{idx}", type="primary"):
                        st.session_state.handout_basket.append(item)
                        st.toast(f"Added {item['file']} (P.{item['page']+1}) to basket!")
                        st.rerun()

                    st.markdown("---")
                    
                    with open(item['path'], "rb") as pdf_f:
                        st.download_button(
                            label="📥 Download Full PDF",
                            data=pdf_f.read(),
                            file_name=item['file'],
                            mime="application/pdf",
                            key=f"dl_t1_{idx}"
                        )


# --- TAB 2: PRACTICAL SEARCH (P2 & P4) ---
with tab2:
    st.header("Search Practical Papers (Paper 2 & Paper 4)")
    st.caption("Variants: Paper 2 (02) | Paper 4 (04)")
    keyword_t2 = st.text_input("Enter Practical Keywords (e.g., 'Mail Merge', 'JavaScript', 'Vector Graphics')", key="t2_kw")

    if st.button("Search Practical Papers", type="primary"):
        if keyword_t2:
            with st.spinner("Scanning Practical PDFs..."):
                keywords = [k.strip() for k in keyword_t2.split(",") if k.strip()]
                practical_variants = ["02", "04"]
                st.session_state.practical_results = search_pdfs(keywords, LOCAL_FOLDERS["practical"], practical_variants)
        else:
            st.warning("Please enter a keyword.")

    if st.session_state.practical_results:
        st.write(f"Found **{len(st.session_state.practical_results)}** matching pages:")
        for idx, item in enumerate(st.session_state.practical_results):
            doc_kind = "📝 Question Paper" if item["type"] == "QP" else "🔑 Marking Scheme"
            label = f"📄 {item['file']} | {doc_kind} | Page {item['page'] + 1}"
            
            with st.expander(label):
                col_img, col_actions = st.columns([3, 1])
                
                with col_img:
                    img_data = render_pdf_page_image(item['path'], item['page'])
                    st.image(img_data, use_container_width=True)

                with col_actions:
                    st.write("### Actions")
                    if st.button("➕ Add to Basket", key=f"add_t2_{idx}", type="primary"):
                        st.session_state.handout_basket.append(item)
                        st.toast(f"Added {item['file']} (P.{item['page']+1}) to basket!")
                        st.rerun()

                    st.markdown("---")

                    with open(item['path'], "rb") as pdf_f:
                        st.download_button(
                            label="📥 Download Full PDF",
                            data=pdf_f.read(),
                            file_name=item['file'],
                            mime="application/pdf",
                            key=f"dl_t2_{idx}"
                        )


# --- TAB 3: HANDOUT BASKET / CART ---
with tab3:
    st.header("Worksheet / Handout Builder")
    if st.session_state.handout_basket:
        st.subheader(f"Selected Question/Answer Pages: {len(st.session_state.handout_basket)}")

        for idx, item in enumerate(list(st.session_state.handout_basket)):
            c1, c2 = st.columns([4, 1])
            c1.write(f"{idx+1}. **{item['file']}** (Page {item['page'] + 1})")
            if c2.button("❌ Remove", key=f"remove_basket_{idx}"):
                st.session_state.handout_basket.pop(idx)
                st.rerun()

        st.markdown("---")
        if st.button("🪄 Export Handout to Word Document", type="primary"):
            with st.spinner("Building custom Word document..."):
                doc = create_custom_word_handout(st.session_state.handout_basket, SYLLABUS_CODE)
                
                target_filename = f"{SYLLABUS_CODE}_IT_Handout.docx"
                doc_buffer = io.BytesIO()
                doc.save(doc_buffer)
                doc_buffer.seek(0)

                st.download_button(
                    label="📥 Click for final Download to Local Drive",
                    data=doc_buffer,
                    file_name=target_filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
    else:
        st.info("Your basket is empty. Add pages from Tab 1 or Tab 2.")


# --- TAB 4: SOURCE FILES (ZIP) ---
with tab4:
    st.header("Download Practical Source Files (ZIP)")
    c1, c2, c3 = st.columns(3)
    with c1:
        z_year = st.selectbox("Select Year", [str(y) for y in range(2026, 2018, -1)])
    with c2:
        z_session = st.selectbox("Select Session", ["March (m)", "June (s)", "Nov (w)"])
        session_code = z_session.split("(")[1].replace(")", "")
    with c3:
        z_paper = st.selectbox("Select Paper Component", ["02 (Paper 2)", "04 (Paper 4)"])
        paper_code = z_paper.split()[0]

    short_year = z_year[-2:]
    expected_zip_name = f"9626_{session_code}{short_year}_sf_{paper_code}.zip"
    zip_path = os.path.join(LOCAL_FOLDERS["zips"], expected_zip_name)

    st.markdown("---")
    if os.path.exists(zip_path):
        st.success(f"Found Source File: `{expected_zip_name}`")
        with open(zip_path, "rb") as zf:
            st.download_button(
                label=f"📦 Download {expected_zip_name}",
                data=zf,
                file_name=expected_zip_name,
                mime="application/zip"
            )
    else:
        st.warning(f"Source file `{expected_zip_name}` is not available locally in `{LOCAL_FOLDERS['zips']}`. Use the Admin Sync button to pull newly uploaded files from Drive.")


# --- TAB 5: ADMIN PANEL ---
with tab5:
    st.header("Admin Panel")

    admin_password = st.secrets.get("ADMIN_PASSWORD")

    if not admin_password:
        st.error(
            "🚨 `ADMIN_PASSWORD` is not configured in your Streamlit Secrets. "
            "Please add `ADMIN_PASSWORD = 'your_password'` to your secrets configuration."
        )
    else:
        pwd = st.text_input("Enter Your Admin Password", type="password")

        if pwd == admin_password:
            st.success("Admin Access Granted")
            
            # --- MANUAL SINGLE FILE UPLOAD ---
            st.subheader("📤 Single File Direct Upload")
            uploaded_file = st.file_uploader("Browse Past Paper PDF or Source File (ZIP)", type=["pdf", "zip"])

            if uploaded_file is not None:
                folder_key, folder_name = determine_target_folder(uploaded_file.name)

                if folder_key is None:
                    st.warning(
                        f"⚠️ Filename `{uploaded_file.name}` does not match Cambridge naming conventions. "
                        "Expected formats: `9626_m19_qp_12.pdf`, `9626_s20_qp_02.pdf`, or `9626_s20_sf_02.zip`."
                    )
                else:
                    st.info(f"🎯 Target Destination Detected: **{folder_name}**")

                    if st.button("🚀 Upload File to Drive & Mirror Locally"):
                        with st.spinner("Processing file..."):
                            file_bytes = uploaded_file.read()

                            # 1. Save locally for instant search access
                            local_dest_dir = LOCAL_FOLDERS[folder_key]
                            local_save_path = os.path.join(local_dest_dir, uploaded_file.name)
                            with open(local_save_path, "wb") as f:
                                f.write(file_bytes)
                            st.info(f"📁 Mirrored file locally to `{local_dest_dir}/{uploaded_file.name}`.")

                            # 2. Upload to Google Drive
                            target_drive_folder_id = FOLDER_IDS[folder_key]
                            drive_result = upload_file_to_drive(
                                file_bytes, 
                                uploaded_file.name, 
                                target_drive_folder_id, 
                                uploaded_file.type
                            )

                            if drive_result:
                                st.success(f"✅ Successfully uploaded `{uploaded_file.name}` to Google Drive!")
                                st.markdown(f"[🔗 View File in Google Drive]({drive_result.get('webViewLink')})")
                            else:
                                st.error("❌ Failed to upload to Google Drive. Check your connection or secrets.")


# ==========================================
# 7. FOOTER
# ==========================================
st.markdown("---")
st.markdown(
    """
    <div style="text-align: center; width: 100%;">
        <p style="font-size: 20px; font-weight: bold; margin-bottom: 5px;">✨ Digital 9626 Information Technology Resource Portal ✨</p>
        <p style="color: gray; font-size: 14px;">Developer: HNHaziqah @ HHartini Computer Science PTES</p>
    </div>
    """,
    unsafe_allow_html=True
)
